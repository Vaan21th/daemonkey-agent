"""
api_routes/intelligence.py · 智识闭环路由 (wish-413999da · phase 1)
==================================================================

11 路由 · UI 直接走 API 操作智识（不走 LLM·节省 token）+ 报告/复盘文件服务:

  POST /verify/claim                  · 佐证按钮 · fact_check.verify_claim
  POST /radar/domains/add             · 直接走 API 加 domain (无 LLM)
  POST /radar/domains/remove          · 直接走 API 删 domain
  POST /radar/feedback                · UI 点 👍/👎/⭐/🗑 落盘
  POST /outcome                       · UI 直接更新机会 outcome 闭环
  POST /favorites                     · UI 点 ⭐ 切换收藏

  GET  /reports/preview/{filename}    · docx 在线预览 (md 源优先 · python-docx 兜底)
  GET  /reports/{filename}            · docx 下载

  GET  /reviews                       · 月度复盘列表 (wish-bf190d9c)
  GET  /reviews/preview/{filename}    · 单份月度复盘预览
  GET  /reviews/file/{filename}       · 单份月度复盘 .md 下载

注: _serve_report_file / _resolve_review_md / _REPORTS_DIR / REVIEWS_DIR
    1:1 copy 旧 closure 实现到 module-level (phase 2 再消除跟 daemon_api
    内 dashboard 路由的代码冗余 · 当前两份实现并存不影响功能)
"""
from __future__ import annotations

from pathlib import Path
from typing import Optional

from fastapi import APIRouter, Body, Header, HTTPException
from fastapi.responses import FileResponse

from api_routes._deps import check_auth


ROOT = Path(__file__).resolve().parent.parent
_REPORTS_DIR = ROOT / "data" / "reports"
REVIEWS_DIR = ROOT / "data" / "reviews"


router = APIRouter()


def _compact_blank_lines(lines: list[str]) -> list[str]:
    """连续空行折叠为一个 · 给 docx 抽取兜底用"""
    out: list[str] = []
    prev_blank = False
    for line in lines:
        is_blank = (not line.strip())
        if is_blank and prev_blank:
            continue
        out.append(line)
        prev_blank = is_blank
    return out


def _serve_report_file(filename: str):
    """安全返回单个 docx · 防越权访问其他目录"""
    if not filename.lower().endswith(".docx"):
        raise HTTPException(400, "only .docx files allowed")
    if "/" in filename or "\\" in filename or ".." in filename:
        raise HTTPException(400, "invalid filename")
    if filename.startswith(".") or filename.startswith("~"):
        raise HTTPException(400, "hidden / temp files forbidden")

    path = (_REPORTS_DIR / filename).resolve()
    try:
        path.relative_to(_REPORTS_DIR.resolve())
    except ValueError:
        raise HTTPException(403, "path escapes reports directory")

    if not path.exists() or not path.is_file():
        raise HTTPException(404, f"report not found: {filename}")

    return FileResponse(
        path,
        media_type=(
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
        filename=filename,
    )


def _resolve_review_md(filename: str) -> "Path":
    """白名单 + 防越权 · 返回安全的 review .md 绝对路径"""
    if not filename.lower().endswith(".md"):
        raise HTTPException(400, "only .md files allowed")
    if "/" in filename or "\\" in filename or ".." in filename:
        raise HTTPException(400, "invalid filename")
    if filename.startswith(".") or filename.startswith("~"):
        raise HTTPException(400, "hidden / temp files forbidden")
    REVIEWS_DIR.mkdir(parents=True, exist_ok=True)
    base = REVIEWS_DIR.resolve()
    path = (base / filename).resolve()
    try:
        path.relative_to(base)
    except ValueError:
        raise HTTPException(403, "path escapes reviews directory")
    if not path.exists() or not path.is_file():
        raise HTTPException(404, f"review not found: {filename}")
    return path


# ─── 卷三十五补丁3 · 智识直通车 (UI → API → workers · 不烧 LLM) ───

@router.post("/verify/claim")
async def post_verify_claim(
    body: dict = Body(...),
    authorization: Optional[str] = Header(None),
):
    """卷三十五补丁3 · 给 UI 的「佐证」按钮 · 验证一条具体 claim

    body: { "claim": "ChatGPT 月活 1000 万", "limit": 5 }
    """
    check_auth(authorization)
    if not isinstance(body, dict):
        raise HTTPException(400, "body 必须是 JSON object")
    claim = (body.get("claim") or "").strip()
    if not claim:
        raise HTTPException(400, "claim 必填")
    limit = int(body.get("limit") or 5)

    from workers.fact_check import verify_claim
    result = verify_claim(claim, limit=limit)
    return result


@router.post("/radar/domains/remove")
async def post_remove_domain(
    body: dict = Body(...),
    authorization: Optional[str] = Header(None),
):
    """卷三十五补丁3 · 用户点 × 删 domain · 直接走 API · 不烧 LLM token

    body:
      { "slug": "<用户自建领域 slug>",
        "sources_action": "reassign|delete|keep",   # 默认 reassign
        "target_domain": "<可选·不传走 fallback 归 self-evolve>"
      }
    """
    check_auth(authorization)
    if not isinstance(body, dict):
        raise HTTPException(400, "body 必须是 JSON object")
    slug = (body.get("slug") or "").strip()
    if not slug:
        raise HTTPException(400, "slug 必填")
    sources_action = (body.get("sources_action") or "reassign").strip()
    target_domain = body.get("target_domain")
    target_domain = target_domain.strip() if isinstance(target_domain, str) and target_domain.strip() else None

    from workers.info_radar import remove_domain as _remove_domain
    result = _remove_domain(
        slug,
        sources_action=sources_action,
        target_domain=target_domain,
    )
    if not result.get("ok"):
        raise HTTPException(400, result.get("error") or "remove_domain 失败")
    return result


@router.post("/radar/domains/add")
async def post_add_domain(
    body: dict = Body(...),
    authorization: Optional[str] = Header(None),
):
    """卷三十五补丁3 · 配套 · 直接走 API 加 domain (无需 LLM)

    body:
      { "slug": "...", "label": "...", "icon": "...",
        "color": "#xxx", "description": "..." }
    """
    check_auth(authorization)
    if not isinstance(body, dict):
        raise HTTPException(400, "body 必须是 JSON object")
    slug = (body.get("slug") or "").strip()
    label = (body.get("label") or "").strip()
    if not slug:
        raise HTTPException(400, "slug 必填")

    from workers.info_radar import add_domain as _add_domain
    try:
        result = _add_domain(
            slug,
            label=label or slug,
            icon=(body.get("icon") or "").strip() or "🧭",
            color=(body.get("color") or "").strip() or "#a0aec0",
            description=(body.get("description") or "").strip(),
        )
    except ValueError as e:
        raise HTTPException(400, str(e))
    return result


@router.post("/radar/feedback")
async def post_radar_feedback(
    body: dict = Body(...),
    authorization: Optional[str] = Header(None),
):
    """卷三十二 · UI 点 👍/👎/⭐/🗑 直接落盘

    body:
      { "item_id": "...",  # md5(url) 前 12
        "feedback": "thumbs_up|thumbs_down|starred|hidden|null",
        "note": "...",     # 可选
        "title_hint": "...", "url_hint": "..."  # 可选 · radar 滚出时兜底
      }
    feedback=null 表示清掉标记
    """
    check_auth(authorization)
    if not isinstance(body, dict):
        raise HTTPException(400, "body 必须是 JSON object")
    item_id = (body.get("item_id") or "").strip()
    if not item_id:
        raise HTTPException(400, "item_id 必填")
    feedback = body.get("feedback")
    from workers.radar_feedback import (
        VALID_FEEDBACK,
        clear_feedback,
        set_feedback,
    )
    if feedback is None or feedback == "":
        result = clear_feedback(item_id)
        return result
    if feedback not in VALID_FEEDBACK:
        raise HTTPException(
            400,
            f"feedback 必须是 {sorted(VALID_FEEDBACK)} 之一·收到 {feedback!r}",
        )
    result = set_feedback(
        item_id,
        feedback,
        note=body.get("note"),
        title_hint=body.get("title_hint"),
        url_hint=body.get("url_hint"),
    )
    if not result.get("ok"):
        raise HTTPException(400, result.get("error") or "set_feedback 失败")
    return result


@router.post("/outcome")
async def post_outcome(
    body: dict = Body(...),
    authorization: Optional[str] = Header(None),
):
    """卷三十一 · UI 直接更新 outcome 闭环 · 不走 LLM/工具循环

    body 形如：
      { "opp_id": "opp-xxxx",
        "status": "abandoned",
        "decision_reason": "因为 ...",
        "actual_revenue_cny": 0, ... }
    """
    check_auth(authorization)
    if not isinstance(body, dict):
        raise HTTPException(400, "body 必须是 JSON object")
    opp_id = (body.get("opp_id") or "").strip()
    if not opp_id:
        raise HTTPException(400, "opp_id 必填")
    from workers.outcomes import record_outcome
    try:
        result = record_outcome(
            opp_id,
            status=body.get("status"),
            decision_reason=body.get("decision_reason"),
            actual_revenue_cny=body.get("actual_revenue_cny"),
            actual_cost_cny=body.get("actual_cost_cny"),
            efficiency_gain=body.get("efficiency_gain"),
            lessons_learned=body.get("lessons_learned"),
            note=body.get("note"),
        )
    except Exception as e:
        raise HTTPException(500, f"record_outcome failed: {e}")
    if not result.get("ok"):
        raise HTTPException(400, result.get("error") or "record 失败")
    return result


@router.post("/favorites")
async def post_favorites(
    body: dict = Body(...),
    authorization: Optional[str] = Header(None),
):
    """卷三十三 · UI 点 ⭐ 切换收藏

    body:
      { "kind": "opportunity|feasibility",
        "ref_id": "opp-xxxx",
        "action": "toggle|remove|add" (默认 toggle),
        "title_hint": "...",     # opportunities/feasibility 标题
        "domain": "...",         # 可选 · 雷达类目
        "note": "..."            # 可选
      }
    雷达条目的 ⭐ 走 /radar/feedback?feedback=starred · 不走这里。
    """
    check_auth(authorization)
    if not isinstance(body, dict):
        raise HTTPException(400, "body 必须是 JSON object")
    kind = (body.get("kind") or "").strip()
    if kind not in ("opportunity", "feasibility"):
        raise HTTPException(
            400,
            f"kind 必须是 opportunity 或 feasibility · 收到 {kind!r}",
        )
    ref_id = (body.get("ref_id") or "").strip()
    if not ref_id:
        raise HTTPException(400, "ref_id 必填")
    action = (body.get("action") or "toggle").lower()
    title_hint = body.get("title_hint") or ""
    domain_hint = body.get("domain") or ""
    note = body.get("note")
    from workers.favorites import (
        add_favorite,
        remove_favorite,
        toggle_favorite,
    )
    if action == "remove":
        r = remove_favorite(kind, ref_id)
    elif action == "add":
        r = add_favorite(
            kind, ref_id,
            title_snap=title_hint, domain=domain_hint, note=note,
        )
    else:
        r = toggle_favorite(
            kind, ref_id,
            title_snap=title_hint, domain=domain_hint, note=note,
        )
    if not r.get("ok"):
        raise HTTPException(400, r.get("error") or "favorites 失败")
    return r


# ─── 卷二十四 + 卷三十三 · 报告库 (docx 预览 / 下载) ───

@router.get("/reports/preview/{filename}")
async def preview_report(
    filename: str,
    authorization: Optional[str] = Header(None),
    token: Optional[str] = None,
):
    """单个报告的在线预览数据 · 卷三十三补丁

    优先：读 `<filename>.md`(新报告生成时同步落的源 · 带 YAML front-matter)
    兜底：用 python-docx 抽取 docx 里的段落 + 标题 + 列表项 · 简陋还原 markdown

    鉴权两条路 (同 download_report):
      1. Authorization: Bearer xxx
      2. query 参数 ?token=xxx
    """
    if token and not authorization:
        authorization = f"Bearer {token}"
    check_auth(authorization)

    if not filename.lower().endswith(".docx"):
        raise HTTPException(400, "filename 必须以 .docx 结尾")
    if "/" in filename or "\\" in filename or ".." in filename:
        raise HTTPException(400, "invalid filename")
    if filename.startswith(".") or filename.startswith("~"):
        raise HTTPException(400, "hidden / temp files forbidden")

    docx_path = (_REPORTS_DIR / filename).resolve()
    try:
        docx_path.relative_to(_REPORTS_DIR.resolve())
    except ValueError:
        raise HTTPException(403, "path escapes reports directory")

    if not docx_path.exists():
        raise HTTPException(404, f"报告不存在: {filename}")

    md_path = docx_path.with_suffix(".md")

    # ─── 方案 1 · 新报告 · 直接读 .md 源 ───
    if md_path.exists():
        raw = md_path.read_text(encoding="utf-8")
        meta: dict = {}
        md_body = raw
        if raw.startswith("---\n"):
            end = raw.find("\n---\n", 4)
            if end > 0:
                fm = raw[4:end]
                md_body = raw[end + 5:].lstrip("\n")
                for line in fm.splitlines():
                    if ":" in line:
                        k, _, v = line.partition(":")
                        meta[k.strip()] = v.strip()
        return {
            "ok": True,
            "name": filename,
            "has_md_source": True,
            "source": "md",
            "markdown": md_body,
            "meta": meta,
        }

    # ─── 方案 2 · 旧报告 · python-docx 抽段落兜底 ───
    try:
        import docx as _docx
    except ImportError:
        raise HTTPException(
            500,
            "python-docx 未安装·无 md 源的旧报告无法预览·只能下载",
        )

    try:
        doc = _docx.Document(str(docx_path))
    except Exception as e:
        raise HTTPException(500, f"docx 解析失败: {type(e).__name__}: {e}")

    lines: list[str] = []
    for para in doc.paragraphs:
        txt = (para.text or "").rstrip()
        if not txt:
            lines.append("")
            continue
        sty = (para.style.name or "").lower() if para.style else ""
        if sty.startswith("heading 1"):
            lines.append(f"# {txt}")
        elif sty.startswith("heading 2"):
            lines.append(f"## {txt}")
        elif sty.startswith("heading 3"):
            lines.append(f"### {txt}")
        elif sty.startswith("heading 4"):
            lines.append(f"#### {txt}")
        elif "list" in sty or sty.startswith("bullet"):
            lines.append(f"- {txt}")
        else:
            lines.append(txt)

    for tbl in doc.tables:
        lines.append("")
        for ri, row in enumerate(tbl.rows):
            cells = ["" if c is None else (c.text or "").strip().replace("\n", " ") for c in row.cells]
            lines.append("| " + " | ".join(cells) + " |")
            if ri == 0:
                lines.append("| " + " | ".join(["---"] * len(cells)) + " |")
        lines.append("")

    md_body = "\n\n".join(_compact_blank_lines(lines))
    title_guess = filename.split("__", 1)[0].replace("_", " ")

    return {
        "ok": True,
        "name": filename,
        "has_md_source": False,
        "source": "docx_extract",
        "markdown": md_body,
        "meta": {"title": title_guess},
        "note": (
            "这份报告是旧版本·没有 markdown 源。当前预览是从 docx 反向抽取的"
            "简陋还原 (标题/段落/列表/表格)。后续生成的报告会自动有 md 源。"
        ),
    }


@router.get("/reports/{filename}")
async def download_report(
    filename: str,
    authorization: Optional[str] = Header(None),
    token: Optional[str] = None,
):
    """单个 docx 报告下载 · WebUI 点条目跳到这

    鉴权两条路 (任一即可 · 方便手机端用 <a href>):
      1. Authorization: Bearer xxx  · 跟 /chat、/dashboard 一致
      2. query 参数 ?token=xxx       · 浏览器 <a> 不能加 header · 走 query
    """
    if token and not authorization:
        authorization = f"Bearer {token}"
    check_auth(authorization)
    return _serve_report_file(filename)


# ─── 卷四十六 II · 月度复盘 reviews (wish-bf190d9c) ───

@router.get("/reviews")
async def list_reviews_endpoint(
    authorization: Optional[str] = Header(None),
    token: Optional[str] = None,
):
    """列 data/reviews/ 下所有月度复盘 · 按 mtime desc。"""
    if token and not authorization:
        authorization = f"Bearer {token}"
    check_auth(authorization)
    try:
        from workers.review_generator import list_reviews
        items = list_reviews()
        return {"ok": True, "count": len(items), "items": items}
    except Exception as e:
        return {"ok": False, "error": f"{type(e).__name__}: {e}", "items": []}


@router.get("/reviews/preview/{filename}")
async def preview_review(
    filename: str,
    authorization: Optional[str] = Header(None),
    token: Optional[str] = None,
):
    """单份月度复盘预览 · 返 markdown + meta。 给 webui mdRender。"""
    if token and not authorization:
        authorization = f"Bearer {token}"
    check_auth(authorization)
    path = _resolve_review_md(filename)
    raw = path.read_text(encoding="utf-8")
    is_final = filename.lower().endswith("-final.md")
    return {
        "ok": True,
        "name": filename,
        "markdown": raw,
        "status": "final" if is_final else "draft",
        "size_bytes": path.stat().st_size,
        "path": str(path.relative_to(ROOT)).replace("\\", "/"),
    }


@router.get("/reviews/file/{filename}")
async def download_review(
    filename: str,
    authorization: Optional[str] = Header(None),
    token: Optional[str] = None,
):
    """单份月度复盘原始 .md 下载 · 浏览器走系统默认应用。"""
    if token and not authorization:
        authorization = f"Bearer {token}"
    check_auth(authorization)
    path = _resolve_review_md(filename)
    return FileResponse(
        path,
        media_type="text/markdown; charset=utf-8",
        filename=filename,
    )
