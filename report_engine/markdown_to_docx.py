"""
report_engine/markdown_to_docx.py
==================================

通用 markdown → docx 渲染器

源流：抽象自一套成熟的 build-doc.py 渲染器。

设计差异：
  - 视觉规范不再 hardcoded · 改为接受 Theme 对象（见 themes.py）
  - 不再 hardcoded 多份 DOCS_CONFIG · 改为接受单文档参数
  - 公共 API `render_report` · 上层只需 (markdown_text, output_path, cover, theme)

支持的 markdown 元素：
  - 标题（1-6 级 · 字号 / 颜色随级别变化）
  - 段落 + 行内格式（**bold** / `code` / *italic*）
  - 无序列表 / 有序列表（多级缩进）
  - 表格（自带表头底色 + 隔行灰底）
  - 引用块（> · 浅底色 + 左侧竖线）
  - 代码块（```fenced``` · 等宽字体 + 灰底）
  - 水平分割线 `---`
  - 图片 `![alt](path.png)` · 路径相对源 markdown 文件

不支持（暂未实现 · 实际生产中很少用）：
  - 链接 `[text](url)` · 行内文字会保留 markdown 原文
  - 行内 HTML
  - 嵌套表格 / 嵌套引用

依赖：
  pip install python-docx
"""
from __future__ import annotations

import re
import sys
import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from .themes import Theme, get_theme, THEME_OPUS_STUDIO


# ──────────────────────────────────────────────────────────
# 文件名容错（docx 被 Word 锁住自动换名）
# ──────────────────────────────────────────────────────────


def resolve_writable_path(target: Path) -> Path:
    """如果默认 docx 被 Word 等占用·自动换名为 -v2 / -v3 ...

    _resolve_out_docx · 保留这个稳定性容错。
    """
    if not target.exists():
        return target
    try:
        with open(target, "a"):
            pass
        return target
    except PermissionError:
        n = 2
        while True:
            alt = target.with_name(f"{target.stem}-v{n}{target.suffix}")
            if not alt.exists():
                return alt
            try:
                with open(alt, "a"):
                    pass
                return alt
            except PermissionError:
                n += 1


# ──────────────────────────────────────────────────────────
# docx 底层 helper（接受 theme · 不再 hardcoded）
# ──────────────────────────────────────────────────────────


def _rgb(t: tuple[int, int, int]) -> RGBColor:
    return RGBColor(*t)


def set_font(
    run,
    theme: Theme,
    *,
    name: Optional[str] = None,
    size: float = 10.5,
    bold: bool = False,
    italic: bool = False,
    color: Optional[RGBColor] = None,
):
    """统一设置 run 字体 · 中英文都设 · 颜色可选"""
    n = name or theme.font_cjk
    run.font.name = n
    run.element.rPr.rFonts.set(qn("w:eastAsia"), n)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, color_hex: str):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_paragraph_shading(paragraph, color_hex: str):
    pPr = paragraph._p.get_or_add_pPr()
    shd = parse_xml(
        f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{color_hex}"/>'
    )
    pPr.append(shd)


def set_paragraph_border(
    paragraph, side: str = "left", size: int = 24, color: str = "2B6CB0"
):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:{side} w:val="single" w:sz="{size}" w:space="6" w:color="{color}"/>'
        "</w:pBdr>"
    )
    pPr.append(pBdr)


# ──────────────────────────────────────────────────────────
# 行内格式解析（**bold** / `code` / *italic*）
# ──────────────────────────────────────────────────────────

INLINE_RE = re.compile(
    r"(\*\*[^*]+?\*\*)"      # bold
    r"|(`[^`]+?`)"           # inline code
    r"|(\*[^*]+?\*)"         # italic
)


def add_inline_runs(
    paragraph,
    text: str,
    theme: Theme,
    *,
    base_size: float = 10.5,
    base_color: Optional[RGBColor] = None,
    base_bold: bool = False,
):
    """处理一段含 inline 格式的文本 · 切分成多个 run 加入 paragraph"""
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            r = paragraph.add_run(text[pos: m.start()])
            set_font(r, theme, size=base_size, color=base_color, bold=base_bold)
        if m.group(1):
            r = paragraph.add_run(m.group(1)[2:-2])
            set_font(r, theme, size=base_size, color=base_color, bold=True)
        elif m.group(2):
            r = paragraph.add_run(m.group(2)[1:-1])
            set_font(r, theme,
                     name=theme.font_en,
                     size=base_size - 0.5,
                     color=_rgb(theme.color_code_inline),
                     bold=base_bold)
        elif m.group(3):
            r = paragraph.add_run(m.group(3)[1:-1])
            set_font(r, theme, size=base_size, color=base_color,
                     bold=base_bold, italic=True)
        pos = m.end()
    if pos < len(text):
        r = paragraph.add_run(text[pos:])
        set_font(r, theme, size=base_size, color=base_color, bold=base_bold)


# ──────────────────────────────────────────────────────────
# 块级渲染（heading / para / list / quote / table / code / image / hr）
# ──────────────────────────────────────────────────────────


def add_heading(doc, text: str, theme: Theme, level: int = 1):
    sizes = {1: 22, 2: 16, 3: 13, 4: 12, 5: 11, 6: 11}
    color_map = {
        1: _rgb(theme.color_title),
        2: _rgb(theme.color_h2),
        3: _rgb(theme.color_h3),
        4: _rgb(theme.color_h3),
        5: _rgb(theme.color_h3),
        6: _rgb(theme.color_h3),
    }
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(14 if level <= 2 else 8)
    h.paragraph_format.space_after = Pt(6)
    add_inline_runs(
        h, text, theme,
        base_size=sizes.get(level, 11),
        base_color=color_map.get(level, _rgb(theme.color_h3)),
        base_bold=True,
    )
    return h


def add_para(doc, text: str, theme: Theme, size: float = 10.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.4
    add_inline_runs(p, text, theme, base_size=size)
    return p


def add_quote(doc, lines: list[str], theme: Theme):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.4)
    set_paragraph_shading(p, theme.quote_fill)
    set_paragraph_border(p, "left", size=24, color=theme.quote_border)
    text = "\n".join(lines)
    add_inline_runs(p, text, theme, base_size=10, base_color=_rgb(theme.color_quote))
    return p


def add_list_item(
    doc, text: str, theme: Theme,
    indent_level: int = 0, ordered: bool = False, number: int = 1,
):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6 + indent_level * 0.6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.35
    bullet = f"{number}. " if ordered else "• "
    r = p.add_run(bullet)
    set_font(r, theme, size=10.5, color=_rgb(theme.color_h2), bold=True)
    add_inline_runs(p, text, theme, base_size=10.5)
    return p


def add_code_block(doc, lines: list[str], theme: Theme):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.line_spacing = 1.2
    set_paragraph_shading(p, theme.code_fill)
    text = "\n".join(lines)
    r = p.add_run(text)
    set_font(r, theme, name=theme.font_en, size=9.5,
             color=RGBColor(0x2D, 0x37, 0x48))


def add_table(doc, headers: list[str], rows: list[list[str]], theme: Theme):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline_runs(p, header, theme,
                        base_size=10,
                        base_color=RGBColor(0xFF, 0xFF, 0xFF),
                        base_bold=True)
        set_cell_shading(cell, theme.table_header_fill)

    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            if c_idx >= len(headers):
                continue
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            add_inline_runs(p, cell_text, theme, base_size=9.5)
            if r_idx % 2 == 1:
                set_cell_shading(cell, theme.table_alt_fill)

    doc.add_paragraph()
    return table


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '<w:bottom w:val="single" w:sz="6" w:space="1" w:color="CBD5E0"/>'
        "</w:pBdr>"
    )
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)


def add_image_with_caption(doc, img_path: Path, alt: str, theme: Theme):
    """嵌入图片 + 居中图注 · 图片不存在时显示占位提示框"""
    if img_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run()
        try:
            r.add_picture(str(img_path), width=Cm(16.5))
        except Exception as e:
            r2 = p.add_run(f"[ 图片加载失败:{img_path.name} ({e}) ]")
            set_font(r2, theme, size=9, color=_rgb(theme.color_hint), italic=True)

        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_before = Pt(0)
        cap.paragraph_format.space_after = Pt(10)
        cr = cap.add_run(f"图  {alt}")
        set_font(cr, theme, size=9, color=_rgb(theme.color_quote), italic=True)
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(10)
        set_paragraph_shading(p, theme.placeholder_fill)
        set_paragraph_border(p, "left", size=24, color=theme.placeholder_border)
        r = p.add_run(
            f"  待插入截图：{alt}\n  （文件路径：{img_path.name}）"
        )
        set_font(r, theme, size=10,
                 color=RGBColor(0xC0, 0x5A, 0x12), italic=True)


# ──────────────────────────────────────────────────────────
# Markdown 主解析器
# ──────────────────────────────────────────────────────────

TABLE_SEP_RE = re.compile(r"^\s*\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$")
HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
ULIST_RE = re.compile(r"^(\s*)[-*+]\s+(.+)$")
OLIST_RE = re.compile(r"^(\s*)(\d+)\.\s+(.+)$")
CODE_FENCE_RE = re.compile(r"^```")
IMAGE_RE = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)")


def split_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def render_markdown_to_doc(
    doc,
    md_text: str,
    theme: Theme,
    *,
    here_dir: Optional[Path] = None,
):
    """主 markdown → docx 渲染器（接受 theme + here_dir for 相对图片路径）"""
    lines = md_text.replace("\r\n", "\n").split("\n")
    i = 0
    n = len(lines)
    here_dir = here_dir or Path.cwd()

    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if CODE_FENCE_RE.match(stripped):
            j = i + 1
            buf: list[str] = []
            while j < n and not CODE_FENCE_RE.match(lines[j].strip()):
                buf.append(lines[j])
                j += 1
            add_code_block(doc, buf, theme)
            i = j + 1
            continue

        if stripped == "---":
            add_horizontal_rule(doc)
            i += 1
            continue

        m_h = HEADING_RE.match(line)
        if m_h:
            level = len(m_h.group(1))
            add_heading(doc, m_h.group(2), theme, level=level)
            i += 1
            continue

        if "|" in stripped and i + 1 < n and TABLE_SEP_RE.match(lines[i + 1]):
            headers = split_table_row(lines[i])
            rows: list[list[str]] = []
            j = i + 2
            while j < n and "|" in lines[j].strip() and lines[j].strip():
                if TABLE_SEP_RE.match(lines[j]):
                    j += 1
                    continue
                rows.append(split_table_row(lines[j]))
                j += 1
            add_table(doc, headers, rows, theme)
            i = j
            continue

        if stripped.startswith(">"):
            buf = []
            j = i
            while j < n and lines[j].lstrip().startswith(">"):
                qt = lines[j].lstrip()[1:].lstrip()
                buf.append(qt)
                j += 1
            add_quote(doc, buf, theme)
            i = j
            continue

        m_ul = ULIST_RE.match(line)
        m_ol = OLIST_RE.match(line)
        if m_ul or m_ol:
            j = i
            while j < n:
                lj = lines[j]
                mu = ULIST_RE.match(lj)
                mo = OLIST_RE.match(lj)
                if mu:
                    indent = len(mu.group(1)) // 2
                    add_list_item(doc, mu.group(2), theme,
                                  indent_level=indent, ordered=False)
                elif mo:
                    indent = len(mo.group(1)) // 2
                    add_list_item(doc, mo.group(3), theme,
                                  indent_level=indent,
                                  ordered=True, number=int(mo.group(2)))
                else:
                    if not lj.strip():
                        j += 1
                        break
                    if lj.startswith("   ") or lj.startswith("\t"):
                        last_p = doc.paragraphs[-1]
                        r = last_p.add_run(" " + lj.strip())
                        set_font(r, theme, size=10.5)
                    else:
                        break
                j += 1
            i = j
            continue

        # 行内图片
        if stripped.startswith("!["):
            mm = IMAGE_RE.match(stripped)
            if mm:
                alt = mm.group(1) or "配图"
                rel_path = mm.group(2).strip()
                img_path = (here_dir / rel_path).resolve()
                add_image_with_caption(doc, img_path, alt, theme)
                i += 1
                continue

        # 普通段落：把连续非空行合到一起
        buf = [stripped]
        j = i + 1
        while j < n:
            nx = lines[j]
            if (not nx.strip()
                or HEADING_RE.match(nx)
                or nx.lstrip().startswith(">")
                or ULIST_RE.match(nx) or OLIST_RE.match(nx)
                or CODE_FENCE_RE.match(nx.strip())
                or nx.strip() == "---"
                or ("|" in nx.strip()
                    and j + 1 < n
                    and TABLE_SEP_RE.match(lines[j + 1]))):
                break
            buf.append(nx.strip())
            j += 1
        add_para(doc, " ".join(buf), theme)
        i = j


# ──────────────────────────────────────────────────────────
# 高层 API · render_report
# ──────────────────────────────────────────────────────────


def _build_cover_page(doc, cover: dict, theme: Theme):
    """渲染封面 · cover 是个 dict · 至少要 title

    可选字段：subtitle / audience / note / footer / date
    """
    for _ in range(5):
        doc.add_paragraph()

    title = cover.get("title", "(未命名报告)")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    set_font(r, theme, size=26, bold=True, color=_rgb(theme.color_title))

    if cover.get("subtitle"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(cover["subtitle"])
        set_font(r, theme, size=18, color=_rgb(theme.color_h2))

    doc.add_paragraph()

    if cover.get("note"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(cover["note"])
        set_font(r, theme, size=12, color=_rgb(theme.color_quote))

    for _ in range(6):
        doc.add_paragraph()

    # footer 区域：日期 + audience + footer 自定义
    date_str = cover.get("date") or datetime.date.today().strftime("%Y 年 %m 月 %d 日")
    parts = [f"文档版本:v1.0  ·  {date_str}"]
    if cover.get("audience"):
        parts.append(cover["audience"])
    if cover.get("footer"):
        parts.append(cover["footer"])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("\n".join(parts))
    set_font(r, theme, size=10, color=_rgb(theme.color_hint))

    doc.add_page_break()


def render_report(
    md_text: str,
    output_path: Path,
    *,
    cover: Optional[dict] = None,
    theme: str | Theme = "opus_studio",
    here_dir: Optional[Path] = None,
    strip_h1_prefix: Optional[str] = None,
) -> Path:
    """高层 API：把 markdown 渲染成 docx · 返回实际写入的路径

    参数：
        md_text         markdown 主稿内容（含 # 一级标题以下的所有内容）
        output_path     目标 docx 路径（文件被占用时自动加 -v2 / -v3）
        cover           封面信息 dict · 字段： title / subtitle / note /
                                              audience / footer / date
                                              · None → 不渲染封面页（直接正文）
        theme           主题名（'opus_studio' / 'manju'）或 Theme 对象
        here_dir        解析 markdown 中相对图片路径的基准目录
                        · 默认 output_path.parent
        strip_h1_prefix 如果 markdown 顶部一级标题已在 cover 体现 · 用此字符串去掉
                        · 例: '# 本周雷达报告' → cover 已经有标题就传 strip_h1_prefix

    返回：
        实际写入的 docx Path（可能跟 output_path 不同 · 文件被占用时）

    异常：
        - python-docx 未安装：ImportError（在 import 时已经抛）
        - md_text 为空：ValueError
    """
    if not md_text or not md_text.strip():
        raise ValueError("md_text 为空 · 没有可渲染的内容")

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    if isinstance(theme, str):
        theme = get_theme(theme)
    elif not isinstance(theme, Theme):
        theme = THEME_OPUS_STUDIO

    here_dir = here_dir or output_path.parent

    doc = Document()

    style = doc.styles["Normal"]
    f = style.font
    f.name = theme.font_cjk
    f.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), theme.font_cjk)

    for section in doc.sections:
        section.top_margin = Cm(2.4)
        section.bottom_margin = Cm(2.4)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    if cover:
        _build_cover_page(doc, cover, theme)

    text = md_text
    if strip_h1_prefix:
        text = re.sub(
            rf"^{re.escape(strip_h1_prefix)}.*?\n",
            "",
            text,
            count=1,
        )

    render_markdown_to_doc(doc, text, theme, here_dir=here_dir)

    final_path = resolve_writable_path(output_path)
    doc.save(final_path)
    return final_path


# ──────────────────────────────────────────────────────────
# CLI（独立调试用）
# ──────────────────────────────────────────────────────────

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("用法: python -m report_engine.markdown_to_docx <src.md> <out.docx> [theme]")
        print("可选 theme: opus_studio (默认) / manju")
        sys.exit(1)
    src = Path(sys.argv[1])
    out = Path(sys.argv[2])
    th = sys.argv[3] if len(sys.argv) > 3 else "opus_studio"

    if not src.exists():
        print(f"[err] 找不到 src markdown: {src}")
        sys.exit(1)

    md = src.read_text(encoding="utf-8")
    cover = {"title": src.stem, "audience": "CLI 单跑生成"}
    final = render_report(
        md, out,
        cover=cover, theme=th, here_dir=src.parent,
    )
    print(f"[ok] 已生成 {final}  ({final.stat().st_size / 1024:.1f} KB)")
